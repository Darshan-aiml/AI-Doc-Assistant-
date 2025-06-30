import gradio as gr
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.text_splitter import CharacterTextSplitter
from langchain.chains import RetrievalQA
from langchain.docstore.document import Document as LCDocument
from langchain.llms import HuggingFacePipeline
from transformers import pipeline
import traceback

qa_system = None

# 📄 Read text from file
def read_document(file_obj):
    try:
        path = file_obj.name
        if path.endswith(".pdf"):
            reader = PdfReader(path)
            return " ".join([p.extract_text() for p in reader.pages if p.extract_text()])
        elif path.endswith(".docx"):
            doc = DocxDocument(path)
            return " ".join([p.text for p in doc.paragraphs])
        elif path.endswith(".txt"):
            with open(path, "r", encoding="utf-8") as f:
                return f.read()
        else:
            return ""
    except Exception as e:
        print("❌ Error reading file:", traceback.format_exc())
        return ""

# 🧠 Create Retrieval QA system
def setup_qa(text):
    splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
    texts = splitter.split_text(text)
    docs = [LCDocument(page_content=t) for t in texts]

    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    db = FAISS.from_documents(docs, embeddings)

    hf_pipe = pipeline("text2text-generation", model="google/flan-t5-base")
    llm = HuggingFacePipeline(pipeline=hf_pipe)

    qa = RetrievalQA.from_chain_type(
        llm=llm,
        retriever=db.as_retriever()
    )
    return qa

# 🔼 Upload and process file
def upload_file(file):
    global qa_system
    try:
        text = read_document(file)
        if not text.strip():
            return "❌ No text found. Try a different file."
        qa_system = setup_qa(text)
        return "✅ Document processed! You can now ask questions."
    except Exception as e:
        return f"❌ Error: {str(e)}"

# 💬 Answer questions
def ask_question(q):
    if qa_system is None:
        return "❌ Please upload a document first."
    try:
        return qa_system.run(q)
    except Exception as e:
        print("❌ Error answering question:", traceback.format_exc())
        return f"❌ Error answering: {str(e)}"

# 🌐 Gradio UI
with gr.Blocks() as demo:
    gr.Markdown("## 📄 AI Document Q&A Assistant ")
    file = gr.File(label="Upload a document", file_types=[".pdf", ".docx", ".txt"])
    status = gr.Textbox(label="Status")
    file.upload(upload_file, inputs=file, outputs=status)

    question = gr.Textbox(label="Ask a question about the document")
    answer = gr.Textbox(label="Answer")
    question.submit(ask_question, inputs=question, outputs=answer)

demo.launch()
